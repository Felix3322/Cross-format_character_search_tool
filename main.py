import sys
import os
import re
import fnmatch
from PyQt6.QtWidgets import (QApplication, QMainWindow, QVBoxLayout, QWidget, QLineEdit, QListWidget,
                             QTextEdit, QPushButton, QHBoxLayout, QLabel, QMessageBox, QDialog, QCheckBox,
                             QComboBox)
from PyQt6.QtSvgWidgets import QSvgWidget
from PyQt6.QtCore import Qt, QUrl, QTimer
from PyQt6.QtGui import QPixmap, QTextCursor, QTextCharFormat, QColor

# 需要安装第三方库
try:
    import docx
    from openpyxl import load_workbook
except ImportError:
    print("请安装 python-docx 和 openpyxl 库以支持 Office 文件格式。")
    sys.exit(1)

class TextSearchApp(QMainWindow):
    def __init__(self):
        super().__init__()

        self.setWindowTitle("文本查找和替换工具")
        self.setGeometry(100, 100, 800, 600)

        # 搜索关键字输入框
        self.search_label = QLabel("请输入要查找的关键字:")
        self.search_input = QLineEdit()

        # 替换关键字输入框
        self.replace_label = QLabel("请输入替换内容:")
        self.replace_input = QLineEdit()

        # 正则表达式复选框
        self.regex_checkbox = QCheckBox("使用正则表达式")

        # 编码格式选择
        self.encoding_label = QLabel("选择编码格式:")
        self.encoding_combo = QComboBox()
        self.encoding_combo.addItems(['utf-8', 'gbk', 'gb2312', 'ascii', 'latin1'])

        # 文件过滤输入框
        self.file_filter_label = QLabel("文件过滤（使用分号分隔，支持通配符，例如 *.txt;*.docx）:")
        self.file_filter_input = QLineEdit("*.txt;*.docx")

        # 显示拖放区域
        self.drop_label = QLabel("将文件夹拖放到此处")
        self.drop_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.drop_label.setStyleSheet("border: 2px dashed gray; padding: 20px;")
        self.drop_label.setAcceptDrops(True)

        # 结果列表框
        self.result_list = QListWidget()

        # 文件预览框
        self.file_preview = QTextEdit()
        self.file_preview.setReadOnly(False)  # 设置为可编辑
        self.file_preview.document().setUndoRedoEnabled(True)

        # 实时保存定时器
        self.save_timer = QTimer()
        self.save_timer.setSingleShot(True)
        self.save_timer.timeout.connect(self.save_current_content)
        self.file_preview.textChanged.connect(self.on_text_changed)

        # 搜索按钮
        self.search_button = QPushButton("开始查找")
        self.search_button.clicked.connect(self.search_files)

        # "下一个匹配项" 按钮
        self.next_match_button = QPushButton("下一个匹配项")
        self.next_match_button.clicked.connect(self.go_to_next_match)

        # "替换当前选中项" 按钮
        self.replace_selection_button = QPushButton("替换当前选中项")
        self.replace_selection_button.clicked.connect(self.replace_current_selection)

        # "替换当前文件所有匹配项" 按钮
        self.replace_current_button = QPushButton("替换当前文件所有匹配项")
        self.replace_current_button.clicked.connect(self.replace_current_file)

        # "替换所有文件匹配项" 按钮
        self.replace_all_button = QPushButton("替换所有文件匹配项")
        self.replace_all_button.clicked.connect(self.replace_all_files)

        # "撤销更改" 按钮
        self.undo_button = QPushButton("撤销更改")
        self.undo_button.clicked.connect(self.undo_last_operation)

        # 布局设置
        layout = QVBoxLayout()
        layout.addWidget(self.search_label)
        layout.addWidget(self.search_input)
        layout.addWidget(self.replace_label)
        layout.addWidget(self.replace_input)
        layout.addWidget(self.regex_checkbox)
        layout.addWidget(self.encoding_label)
        layout.addWidget(self.encoding_combo)
        layout.addWidget(self.file_filter_label)
        layout.addWidget(self.file_filter_input)
        layout.addWidget(self.drop_label)
        layout.addWidget(self.search_button)
        layout.addWidget(self.result_list)

        button_layout = QHBoxLayout()
        button_layout.addWidget(self.next_match_button)
        button_layout.addWidget(self.replace_selection_button)
        button_layout.addWidget(self.replace_current_button)
        button_layout.addWidget(self.replace_all_button)
        button_layout.addWidget(self.undo_button)
        layout.addLayout(button_layout)

        layout.addWidget(self.file_preview)

        # 主窗口设置
        container = QWidget()
        container.setLayout(layout)
        self.setCentralWidget(container)

        # 在右上角显示 "i" 图标
        self.create_info_icon()

        # 设置拖放功能
        self.setAcceptDrops(True)
        self.folder_path = ""
        self.matches = []  # 保存所有匹配项的位置
        self.current_match_index = -1  # 当前的匹配项索引

        self.undo_stack = []  # 撤销栈

        # 结果点击事件
        self.result_list.itemClicked.connect(self.preview_file)

    def create_info_icon(self):
        """在右上角显示 i 信息图标"""
        self.info_label = QLabel(self)
        self.info_label.setText("ℹ️")  # 使用 i 图标
        self.info_label.setStyleSheet("font-size: 24px; cursor: pointer;")
        self.info_label.setToolTip("本软件遵循 CC BY-NC 2.0 许可证")
        self.info_label.move(self.width() - 30, 10)  # 右上角对齐
        self.info_label.mousePressEvent = self.show_license_info

    def show_license_info(self, event):
        """显示带有 CC Logo、BY、NC 图标的详情页面"""
        dialog = QDialog(self)
        dialog.setWindowTitle("软件协议")

        # 布局
        dialog_layout = QVBoxLayout()

        # CC 图标的 SVG 数据（省略，保持不变）
        # ...（保持原样）

        dialog.setLayout(dialog_layout)
        dialog.exec()

    def get_search_pattern(self, text):
        if self.regex_checkbox.isChecked():
            try:
                pattern = re.compile(text)
            except re.error:
                self.show_error_message("无效的正则表达式！")
                return None
        else:
            pattern = re.compile(re.escape(text))
        return pattern

    def search_files(self):
        self.result_list.clear()
        self.file_preview.clear()
        self.matches.clear()
        self.current_match_index = -1

        search_term = self.search_input.text()
        if not search_term or not self.folder_path:
            self.show_error_message("请输入关键字并拖放文件夹！")
            return

        file_filters = self.file_filter_input.text().split(';')
        file_filters = [f.strip() for f in file_filters if f.strip()]
        encoding = self.encoding_combo.currentText()

        pattern = self.get_search_pattern(search_term)
        if not pattern:
            return

        # 遍历文件夹中的所有文件
        for root, dirs, files in os.walk(self.folder_path):
            for file_name in files:
                # 文件过滤
                matched = False
                for file_filter in file_filters:
                    if fnmatch.fnmatch(file_name, file_filter):
                        matched = True
                        break
                if not matched:
                    continue

                file_path = os.path.join(root, file_name)
                try:
                    if file_name.endswith('.docx'):
                        content = self.read_docx(file_path)
                    elif file_name.endswith('.xlsx'):
                        content = self.read_xlsx(file_path)
                    else:
                        with open(file_path, 'r', encoding=encoding, errors='ignore') as file:
                            content = file.read()

                    matches = list(pattern.finditer(content))

                    if matches:
                        self.result_list.addItem(f"{file_path} - {len(matches)} 处匹配")
                except Exception as e:
                    self.show_error_message(f"无法读取文件: {file_path}\n错误信息: {e}")

    def read_docx(self, file_path):
        doc = docx.Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return '\n'.join(full_text)

    def read_xlsx(self, file_path):
        wb = load_workbook(file_path)
        full_text = []
        for sheet in wb:
            for row in sheet.iter_rows(values_only=True):
                full_text.append(' '.join([str(cell) if cell is not None else '' for cell in row]))
        return '\n'.join(full_text)

    def preview_file(self, item):
        file_info = item.text().split(" - ")[0]
        search_term = self.search_input.text()
        encoding = self.encoding_combo.currentText()

        try:
            if file_info.endswith('.docx'):
                content = self.read_docx(file_info)
                self.current_file_type = 'docx'
            elif file_info.endswith('.xlsx'):
                content = self.read_xlsx(file_info)
                self.current_file_type = 'xlsx'
            else:
                with open(file_info, 'r', encoding=encoding, errors='ignore') as file:
                    content = file.read()
                self.current_file_type = 'text'

            self.file_preview.setPlainText(content)
            self.current_file_path = file_info  # 保存当前文件路径

            self.matches.clear()
            self.current_match_index = -1

            pattern = self.get_search_pattern(search_term)
            if not pattern:
                return

            self.file_preview.moveCursor(QTextCursor.MoveOperation.Start)
            highlight_format = QTextCharFormat()
            highlight_format.setBackground(QColor("yellow"))
            highlight_format.setForeground(QColor("black"))

            # 移除之前的高亮
            self.file_preview.setExtraSelections([])

            extraSelections = []

            for match in pattern.finditer(content):
                start_pos = match.start()
                end_pos = match.end()
                selection = QTextEdit.ExtraSelection()
                selection.cursor = self.file_preview.textCursor()
                selection.cursor.setPosition(start_pos)
                selection.cursor.setPosition(end_pos, QTextCursor.MoveMode.KeepAnchor)
                selection.format = highlight_format
                extraSelections.append(selection)

                self.matches.append((start_pos, end_pos))

            self.file_preview.setExtraSelections(extraSelections)

            if self.matches:
                self.current_match_index = 0
                self.go_to_match(self.current_match_index)

        except Exception as e:
            self.show_error_message(f"无法预览文件: {file_info}\n错误信息: {e}")

    def go_to_match(self, index):
        if index < 0 or index >= len(self.matches):
            return

        start_pos, end_pos = self.matches[index]
        cursor = self.file_preview.textCursor()
        cursor.setPosition(start_pos)
        cursor.setPosition(end_pos, QTextCursor.MoveMode.KeepAnchor)
        self.file_preview.setTextCursor(cursor)
        self.file_preview.ensureCursorVisible()

    def go_to_next_match(self):
        if self.matches:
            self.current_match_index = (self.current_match_index + 1) % len(self.matches)
            self.go_to_match(self.current_match_index)

    def replace_current_selection(self):
        cursor = self.file_preview.textCursor()
        if cursor.hasSelection():
            selected_text = cursor.selectedText()
            search_term = self.search_input.text()
            replace_term = self.replace_input.text()
            pattern = self.get_search_pattern(search_term)
            if not pattern:
                return
            replaced_text = pattern.sub(replace_term, selected_text)
            cursor.insertText(replaced_text)

    def replace_current_file(self):
        if not self.matches:
            return

        replace_term = self.replace_input.text()
        search_term = self.search_input.text()
        encoding = self.encoding_combo.currentText()

        content = self.file_preview.toPlainText()

        pattern = self.get_search_pattern(search_term)
        if not pattern:
            return

        # 保存原始内容以便撤销
        original_content = content

        new_content, num_subs = pattern.subn(replace_term, content)

        if num_subs > 0:
            # 将操作信息压入撤销栈
            self.undo_stack.append({
                'type': 'replace_current_file',
                'file_path': self.current_file_path,
                'original_content': original_content,
                'num_replacements': num_subs,
                'file_type': self.current_file_type
            })

            self.file_preview.setPlainText(new_content)
            self.save_file(self.current_file_path, new_content, encoding, self.current_file_type)
            self.preview_file(self.result_list.currentItem())

    def replace_all_files(self):
        search_term = self.search_input.text()
        replace_term = self.replace_input.text()
        encoding = self.encoding_combo.currentText()

        if not search_term or not self.folder_path:
            self.show_error_message("请输入关键字并拖放文件夹！")
            return

        file_filters = self.file_filter_input.text().split(';')
        file_filters = [f.strip() for f in file_filters if f.strip()]

        pattern = self.get_search_pattern(search_term)
        if not pattern:
            return

        # 存储被修改的文件信息以便撤销
        modified_files = []

        # 遍历文件夹中的所有文件
        for root, dirs, files in os.walk(self.folder_path):
            for file_name in files:
                # 文件过滤
                matched = False
                for file_filter in file_filters:
                    if fnmatch.fnmatch(file_name, file_filter):
                        matched = True
                        break
                if not matched:
                    continue

                file_path = os.path.join(root, file_name)
                try:
                    if file_name.endswith('.docx'):
                        content = self.read_docx(file_path)
                        file_type = 'docx'
                    elif file_name.endswith('.xlsx'):
                        content = self.read_xlsx(file_path)
                        file_type = 'xlsx'
                    else:
                        with open(file_path, 'r', encoding=encoding, errors='ignore') as file:
                            content = file.read()
                        file_type = 'text'

                    new_content, num_subs = pattern.subn(replace_term, content)

                    if num_subs > 0:
                        # 保存原始内容
                        modified_files.append({
                            'file_path': file_path,
                            'original_content': content,
                            'num_replacements': num_subs,
                            'file_type': file_type
                        })
                        self.save_file(file_path, new_content, encoding, file_type)
                except Exception as e:
                    self.show_error_message(f"无法读取或替换文件: {file_path}\n错误信息: {e}")

        if modified_files:
            # 将操作信息压入撤销栈
            self.undo_stack.append({
                'type': 'replace_all_files',
                'modified_files': modified_files
            })
            self.show_info_message("替换完成！")

    def undo_last_operation(self):
        if not self.undo_stack:
            self.show_info_message("没有可以撤销的操作！")
            return

        last_operation = self.undo_stack.pop()

        encoding = self.encoding_combo.currentText()

        if last_operation['type'] == 'replace_current_file':
            file_path = last_operation['file_path']
            original_content = last_operation['original_content']
            num_replacements = last_operation['num_replacements']
            file_type = last_operation['file_type']

            # 还原文件内容
            self.save_file(file_path, original_content, encoding, file_type)
            self.show_info_message(f"已撤销对文件 {file_path} 的替换，撤销了 {num_replacements} 处替换。")

            # 更新结果列表框
            self.result_list.clear()
            self.result_list.addItem(f"{file_path} - 撤销了 {num_replacements} 处替换")

            # 重新预览文件
            self.preview_file(self.result_list.currentItem())

        elif last_operation['type'] == 'replace_all_files':
            modified_files = last_operation['modified_files']
            self.result_list.clear()

            for file_info in modified_files:
                file_path = file_info['file_path']
                original_content = file_info['original_content']
                num_replacements = file_info['num_replacements']
                file_type = file_info['file_type']

                # 还原文件内容
                self.save_file(file_path, original_content, encoding, file_type)

                # 在结果列表框中显示撤销信息
                self.result_list.addItem(f"{file_path} - 撤销了 {num_replacements} 处替换")

            self.show_info_message("已撤销替换所有文件的操作。")

    def save_file(self, file_path, content, encoding, file_type):
        try:
            if file_type == 'docx':
                self.write_docx(file_path, content)
            elif file_type == 'xlsx':
                self.write_xlsx(file_path, content)
            else:
                with open(file_path, 'w', encoding=encoding, errors='ignore') as file:
                    file.write(content)
        except Exception as e:
            self.show_error_message(f"无法保存文件: {file_path}\n错误信息: {e}")

    def write_docx(self, file_path, content):
        # 由于文本内容已经被修改为纯文本，需要将其写回到 docx 文件中
        doc = docx.Document()
        for line in content.split('\n'):
            doc.add_paragraph(line)
        doc.save(file_path)

    def write_xlsx(self, file_path, content):
        # 由于文本内容已经被修改为纯文本，需要将其写回到 xlsx 文件中
        wb = load_workbook(file_path)
        sheet = wb.active
        for idx, line in enumerate(content.split('\n')):
            # 简单处理，只修改第一列的数据
            sheet.cell(row=idx+1, column=1, value=line)
        wb.save(file_path)

    def on_text_changed(self):
        # 重启定时器，每次文本改变后等待1秒再保存
        self.save_timer.start(1000)

    def save_current_content(self):
        if hasattr(self, 'current_file_path'):
            current_content = self.file_preview.toPlainText()
            encoding = self.encoding_combo.currentText()
            self.save_file(self.current_file_path, current_content, encoding, self.current_file_type)

    def show_error_message(self, message):
        msg_box = QMessageBox(self)
        msg_box.setIcon(QMessageBox.Icon.Critical)
        msg_box.setText(message)
        msg_box.setWindowTitle("错误")
        msg_box.exec()

    def show_info_message(self, message):
        msg_box = QMessageBox(self)
        msg_box.setIcon(QMessageBox.Icon.Information)
        msg_box.setText(message)
        msg_box.setWindowTitle("信息")
        msg_box.exec()

    def dragEnterEvent(self, event):
        if event.mimeData().hasUrls():
            event.acceptProposedAction()

    def dropEvent(self, event):
        urls = event.mimeData().urls()
        if urls:
            path = urls[0].toLocalFile()
            if os.path.isdir(path):
                self.folder_path = path
                self.drop_label.setText(f"已选择文件夹: {path}")
            else:
                self.show_error_message("请拖放一个文件夹！")

if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = TextSearchApp()
    window.show()
    sys.exit(app.exec())
